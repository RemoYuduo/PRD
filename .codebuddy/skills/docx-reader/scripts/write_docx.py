#!/usr/bin/env python3
"""
创建和编写 Word (.docx) 文件

依赖: pip install python-docx
用法: python write_docx.py <output_path> --content <content_file_or_string> [--template <template_path>]

支持从JSON文件读取结构化内容，或从Markdown文件转换
"""

import sys
import argparse
import json
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 请先安装 python-docx 库")
    print("运行: pip install python-docx")
    sys.exit(1)


class DocxWriter:
    """Word文档写入器"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        初始化文档写入器
        
        Args:
            template_path: 可选的模板文件路径
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        
        self._setup_chinese_font()
    
    def _setup_chinese_font(self):
        """设置中文字体支持"""
        try:
            style = self.doc.styles['Normal']
            style.font.name = '微软雅黑'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except:
            pass  # 如果设置失败，使用默认字体
    
    def set_cell_shading(self, cell, color: str):
        """设置单元格背景色"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def add_title(self, text: str, level: int = 0):
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别 (0=文档标题, 1-6=章节标题)
        """
        if level == 0:
            # 文档主标题
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            self.doc.add_heading(text, level=min(level, 9))
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False, 
                      alignment: str = 'left', font_size: int = None,
                      color: tuple = None):
        """
        添加段落
        
        Args:
            text: 段落文本
            bold: 是否加粗
            italic: 是否斜体
            alignment: 对齐方式 (left/center/right/justify)
            font_size: 字体大小(磅)
            color: RGB颜色元组，如 (255, 0, 0)
        """
        para = self.doc.add_paragraph()
        
        # 设置对齐
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        
        if font_size:
            run.font.size = Pt(font_size)
        
        if color:
            run.font.color.rgb = RGBColor(*color)
        
        return para
    
    def add_table(self, headers: List[str], rows: List[List[str]], 
                  header_color: str = '4472C4', style: str = 'Table Grid'):
        """
        添加表格
        
        Args:
            headers: 表头列表
            rows: 数据行列表
            header_color: 表头背景色(十六进制)
            style: 表格样式
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = style
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self.set_cell_shading(header_cells[i], header_color)
        
        # 填充数据
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(cell_data)
        
        return table
    
    def add_key_value_table(self, data: List[Dict[str, str]], 
                            label_color: str = 'E7E6E6'):
        """
        添加键值对表格（如文档信息表）
        
        Args:
            data: 键值对列表，每项为 {'key': '标签', 'value': '值'}
            label_color: 标签列背景色
        """
        # 计算需要的行数（每行2组键值对）
        items_per_row = 2
        num_rows = (len(data) + items_per_row - 1) // items_per_row
        
        table = self.doc.add_table(rows=num_rows, cols=items_per_row * 2)
        table.style = 'Table Grid'
        
        for idx, item in enumerate(data):
            row_idx = idx // items_per_row
            col_offset = (idx % items_per_row) * 2
            
            # 标签单元格
            key_cell = table.rows[row_idx].cells[col_offset]
            key_cell.text = item.get('key', '')
            for para in key_cell.paragraphs:
                for run in para.runs:
                    run.bold = True
            self.set_cell_shading(key_cell, label_color)
            
            # 值单元格
            value_cell = table.rows[row_idx].cells[col_offset + 1]
            value_cell.text = item.get('value', '')
        
        return table
    
    def add_list(self, items: List[str], ordered: bool = False):
        """
        添加列表
        
        Args:
            items: 列表项
            ordered: 是否为有序列表
        """
        for i, item in enumerate(items):
            if ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "• "
            self.doc.add_paragraph(prefix + item)
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_empty_lines(self, count: int = 1):
        """添加空行"""
        for _ in range(count):
            self.doc.add_paragraph()
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
        print(f"文档已保存: {output_path}")
    
    def from_json(self, json_content: Dict[str, Any]):
        """
        从JSON结构创建文档
        
        JSON结构示例:
        {
            "elements": [
                {"type": "title", "text": "文档标题", "level": 0},
                {"type": "heading", "text": "章节标题", "level": 1},
                {"type": "paragraph", "text": "段落内容", "bold": false},
                {"type": "table", "headers": ["列1", "列2"], "rows": [["数据1", "数据2"]]},
                {"type": "list", "items": ["项目1", "项目2"], "ordered": false},
                {"type": "page_break"},
                {"type": "empty_lines", "count": 2}
            ]
        }
        """
        elements = json_content.get('elements', [])
        
        for elem in elements:
            elem_type = elem.get('type', '')
            
            if elem_type == 'title':
                self.add_title(elem.get('text', ''), elem.get('level', 0))
            
            elif elem_type == 'heading':
                self.add_title(elem.get('text', ''), elem.get('level', 1))
            
            elif elem_type == 'paragraph':
                self.add_paragraph(
                    elem.get('text', ''),
                    bold=elem.get('bold', False),
                    italic=elem.get('italic', False),
                    alignment=elem.get('alignment', 'left'),
                    font_size=elem.get('font_size'),
                    color=tuple(elem['color']) if elem.get('color') else None
                )
            
            elif elem_type == 'table':
                self.add_table(
                    elem.get('headers', []),
                    elem.get('rows', []),
                    header_color=elem.get('header_color', '4472C4')
                )
            
            elif elem_type == 'key_value_table':
                self.add_key_value_table(elem.get('data', []))
            
            elif elem_type == 'list':
                self.add_list(
                    elem.get('items', []),
                    ordered=elem.get('ordered', False)
                )
            
            elif elem_type == 'page_break':
                self.add_page_break()
            
            elif elem_type == 'empty_lines':
                self.add_empty_lines(elem.get('count', 1))
    
    def from_markdown(self, md_content: str):
        """
        从Markdown内容创建文档
        
        支持的Markdown语法:
        - # 到 ###### 标题
        - 普通段落
        - - 或 * 无序列表
        - 1. 有序列表
        - | 表格 |
        """
        lines = md_content.strip().split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # 标题
            if line.startswith('#'):
                level = 0
                while level < len(line) and line[level] == '#':
                    level += 1
                text = line[level:].strip()
                if level == 1:
                    self.add_title(text, level=1)
                else:
                    self.add_title(text, level=level)
                i += 1
            
            # 表格
            elif line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 2:
                    # 解析表格
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    rows = []
                    for tl in table_lines[2:]:  # 跳过分隔行
                        if '---' not in tl:
                            row = [cell.strip() for cell in tl.split('|')[1:-1]]
                            rows.append(row)
                    if headers:
                        self.add_table(headers, rows)
            
            # 无序列表
            elif line.startswith('- ') or line.startswith('* '):
                items = []
                while i < len(lines):
                    l = lines[i].strip()
                    if l.startswith('- ') or l.startswith('* '):
                        items.append(l[2:])
                        i += 1
                    else:
                        break
                self.add_list(items, ordered=False)
            
            # 有序列表
            elif line[0].isdigit() and '. ' in line:
                items = []
                while i < len(lines):
                    l = lines[i].strip()
                    if l and l[0].isdigit() and '. ' in l:
                        items.append(l.split('. ', 1)[1])
                        i += 1
                    else:
                        break
                self.add_list(items, ordered=True)
            
            # 普通段落
            else:
                self.add_paragraph(line)
                i += 1


def main():
    parser = argparse.ArgumentParser(description='创建Word文档(.docx)')
    parser.add_argument('output', help='输出文件路径')
    parser.add_argument('--content', '-c', help='内容文件路径(JSON或Markdown)或直接内容')
    parser.add_argument('--template', '-t', help='模板文件路径')
    parser.add_argument('--format', '-f', choices=['json', 'markdown', 'md', 'auto'],
                        default='auto', help='内容格式 (默认: auto自动检测)')
    
    args = parser.parse_args()
    
    writer = DocxWriter(args.template)
    
    if args.content:
        content_path = Path(args.content)
        
        if content_path.exists():
            # 从文件读取
            with open(content_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 自动检测格式
            if args.format == 'auto':
                if content_path.suffix.lower() == '.json':
                    args.format = 'json'
                elif content_path.suffix.lower() in ['.md', '.markdown']:
                    args.format = 'markdown'
                else:
                    # 尝试解析JSON
                    try:
                        json.loads(content)
                        args.format = 'json'
                    except:
                        args.format = 'markdown'
        else:
            # 直接使用内容字符串
            content = args.content
            if args.format == 'auto':
                try:
                    json.loads(content)
                    args.format = 'json'
                except:
                    args.format = 'markdown'
        
        if args.format == 'json':
            json_content = json.loads(content)
            writer.from_json(json_content)
        else:
            writer.from_markdown(content)
    
    output_path = Path(args.output)
    if not output_path.suffix.lower() == '.docx':
        output_path = output_path.with_suffix('.docx')
    
    writer.save(str(output_path))


if __name__ == '__main__':
    main()
